import os
import sys
from datetime import date

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".deps"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = os.path.join(os.path.dirname(__file__), "Plano_Implementacao_25_Modulos_Orbit_Projects.docx")
REPO_URL = "https://github.com/araujofran/sistemaGestaoPro.git"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "6757D9")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(color)
    props.append(underline)
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_status_table(doc, current, gaps, actions, validation):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    labels = [
        ("Situação atual", "EAF5EE", (32, 133, 110)),
        ("O que falta", "FDEEEE", (200, 76, 76)),
        ("Como implementar", "EEEBFF", (103, 87, 217)),
        ("Como validar", "EEF4FB", (52, 120, 201)),
    ]
    values = [current, gaps, actions, validation]
    for row, (label, fill, color), value in zip(table.rows, labels, values):
        row.cells[0].width = Cm(3.6)
        row.cells[1].width = Cm(13.1)
        shade(row.cells[0], fill)
        set_cell_text(row.cells[0], label, bold=True, color=color, size=9)
        set_cell_text(row.cells[1], value, size=9)
    doc.add_paragraph()


modules = [
    (
        "1. Gestão de Projetos",
        "Criação, edição, arquivamento, restauração e exclusão; Scrum, Kanban e híbrido; templates iniciais; clonagem com tarefas, sprints e releases; liderança e modo de gestão.",
        "Biblioteca administrável de templates, configurações específicas por projeto, importação/exportação, governança e permissões por projeto.",
        "Persistir templates no PostgreSQL; criar tabelas projects, project_templates e project_members; expor APIs CRUD; adicionar tela de configurações e políticas por projeto.",
        "Criar projetos nos três modelos; clonar projeto com relacionamentos; arquivar/restaurar; testar acesso permitido e negado por papel; validar exclusão transacional.",
    ),
    (
        "2. Gestão de Tarefas (Issues)",
        "Criação, edição e exclusão; tarefas, histórias, bugs, épicos, incidentes e solicitações; subtarefas; clonagem; mudança entre projetos; campos básicos.",
        "Tipos customizados persistentes, vínculos tipados, hierarquia configurável, campos personalizados, validações obrigatórias, anexos e histórico completo.",
        "Modelar issue_types, issue_links, custom_fields, custom_field_values e attachments; criar serviço de numeração por projeto; adicionar validação de esquema no backend.",
        "Testar cada tipo; criar hierarquia épico→história→subtarefa; mover e clonar preservando dados permitidos; validar obrigatoriedade e anexos; conferir histórico.",
    ),
    (
        "3. Backlog",
        "Backlog por projeto, filtros, estimativas, planejamento de sprint e drag-and-drop entre backlog e sprint.",
        "Ordenação persistente, refinamento estruturado, critérios de pronto, priorização avançada e edição em massa.",
        "Adicionar rank fracionário por projeto; endpoint de reordenação; ações em lote; campos de refinamento e Definition of Ready.",
        "Reordenar centenas de itens sem perder posição; mover em lote; recarregar e confirmar persistência; validar filtros e concorrência.",
    ),
    (
        "4. Scrum",
        "Sprints com nome, meta, datas, status, backlog e pontos; status pode ser planejado, ativo ou concluído; velocidade básica.",
        "Fluxos formais Start/Complete, escopo congelado, capacidade, burndown, burnup, relatório de sprint e tratamento de itens não concluídos.",
        "Criar eventos de sprint e snapshots diários; endpoints start/complete; tabela sprint_capacity; consultas históricas para gráficos.",
        "Iniciar apenas uma sprint por projeto; concluir movendo pendências; validar gráficos contra snapshots; conferir velocidade e capacidade por pessoa.",
    ),
    (
        "5. Kanban",
        "Board Kanban, movimentação por status e limites WIP visíveis.",
        "Colunas customizadas, bloqueio ou alerta real de WIP, swimlanes, políticas de fluxo, lead time e cycle time calculados.",
        "Modelar workflows/status por projeto; registrar timestamps de transição; criar políticas WIP; consultas de tempo de fluxo e swimlanes configuráveis.",
        "Configurar colunas; exceder WIP e conferir regra; medir lead/cycle time em casos conhecidos; validar swimlanes e histórico de transições.",
    ),
    (
        "6. Boards",
        "Boards básicos para projetos Scrum e Kanban, filtros por texto, responsável e prioridade.",
        "Boards independentes, privados/compartilhados, quick filters, cores, layout de cartões e swimlanes.",
        "Criar tabelas boards, board_members, board_filters e board_layout; editor visual; política de visibilidade por equipe.",
        "Criar board privado e compartilhado; validar acesso; salvar filtros/layout; conferir cartões e swimlanes em desktop e mobile.",
    ),
    (
        "7. Planejamento",
        "Roadmap básico por épicos e visualização de releases.",
        "Dependências, marcos, cronograma editável, forecast, caminhos críticos, planejamento entre projetos e Advanced Roadmaps.",
        "Criar modelos milestones e dependencies; componente de timeline; algoritmo inicial de forecast baseado em capacidade e velocidade.",
        "Criar dependências válidas e ciclos inválidos; deslocar datas; validar marcos, forecast e impacto cruzado entre projetos.",
    ),
    (
        "8. Gestão de Releases",
        "Releases demonstrativas com data, status e progresso; clonagem junto ao projeto.",
        "CRUD completo, versionamento, escopo de entrega, notas, aprovação e integração com deployments.",
        "Criar releases, release_issues e release_notes; calcular progresso; fluxo draft→ready→released; API e tela de edição.",
        "Planejar release, associar itens, gerar notas, alterar status e conferir progresso; impedir release inconsistente.",
    ),
    (
        "9. Relatórios",
        "KPIs gerais, visão por responsável, gráfico de velocidade básico e indicadores de fluxo demonstrativos.",
        "Burndown, burnup, sprint/epic/version reports, control chart, CFD, time tracking e workload com dados históricos reais.",
        "Criar camada analítica com snapshots/eventos; consultas agregadas; filtros por período/projeto/equipe; exportação CSV/PDF.",
        "Comparar relatórios com conjunto de dados controlado; validar fusos e períodos; exportar; verificar desempenho com alto volume.",
    ),
    (
        "10. Dashboards",
        "Dashboard geral e KPIs recalculados após alterações.",
        "Dashboards pessoais e compartilhados, gadgets, layout customizável, indicadores customizados e atualização em tempo real.",
        "Modelar dashboards/widgets; editor de grade; WebSocket/SSE; catálogo de métricas e permissões de compartilhamento.",
        "Criar layouts diferentes por usuário; compartilhar; validar atualização ao vivo e autorização de cada gadget.",
    ),
    (
        "11. Busca e Consultas",
        "Pesquisa simples e filtros salvos básicos.",
        "Pesquisa avançada, linguagem semelhante a JQL, filtros compartilhados, assinatura e indexação.",
        "Definir gramática de consulta; parser seguro; índices PostgreSQL full-text/trigram; scheduler de assinaturas.",
        "Testar operadores, datas e combinações; rejeitar consultas inválidas/inseguras; medir desempenho e envio das assinaturas.",
    ),
    (
        "12. Automação",
        "Cadastro inicial de regras na Central, ainda sem motor de execução.",
        "Gatilhos, condições, ações, webhooks, aprovações, escalonamentos, logs e prevenção de loops.",
        "Criar motor orientado a eventos, fila de jobs, worker, catálogo de ações, tentativas e audit trail.",
        "Executar regras em eventos reais; testar condição falsa/verdadeira; retry; idempotência; limite de ciclos e logs.",
    ),
    (
        "13. Workflows",
        "Estados fixos e movimentação básica.",
        "Editor visual, estados por projeto, transições, validações, pós-funções, aprovações e regras de negócio.",
        "Modelar workflow_definitions, states e transitions; máquina de estados no backend; editor visual e publicação versionada.",
        "Publicar workflow; testar transições permitidas/negadas; validar aprovação; migrar itens entre versões sem perda.",
    ),
    (
        "14. Colaboração",
        "Comentários, exclusão de comentários e histórico básico de atividades.",
        "Menções, watchers, compartilhamento, notificações reais, anexos, edição de comentários e preferências.",
        "Criar mentions/watchers/notifications; serviço de e-mail/Teams; central de notificações e preferências por usuário.",
        "Mencionar usuário; acompanhar item; validar notificações e opt-out; checar autorização de anexos e compartilhamento.",
    ),
    (
        "15. Gestão de Tempo",
        "Worklogs, horas gastas, exclusão e total básico por pessoa.",
        "Estimativa original/restante, calendários, aprovação, timesheet e relatórios de esforço.",
        "Modelar worklogs normalizados, estimativas e períodos; tela de timesheet; aprovação opcional e exportação.",
        "Registrar/editar/aprovar horas; validar totais, períodos bloqueados, fusos, exportação e permissões.",
    ),
    (
        "16. Gestão de Equipes",
        "Cadastro, edição e remoção de pessoas; responsáveis, líderes e papéis iniciais.",
        "Contas reais, grupos, times, perfis, convites, disponibilidade e matriz completa de permissões.",
        "Criar users, memberships, groups, teams e role_assignments; telas administrativas e convites com expiração.",
        "Convidar/desativar usuário; testar papéis; validar reatribuição de trabalho e acesso por time/projeto.",
    ),
    (
        "17. Segurança",
        "Proteção básica contra sobrescrita concorrente; ainda sem autenticação.",
        "Login, sessões, RBAC, permissões granulares, auditoria, logs, SSO, OAuth, SAML, proteção CSRF/rate limit e gestão de segredos.",
        "Adotar OIDC (preferencialmente Microsoft Entra ou Google Workspace); cookies seguros; RBAC; audit_log imutável; secrets fora do código.",
        "Testes de acesso horizontal/vertical; expiração de sessão; revisão OWASP; auditoria; restore de backup e rotação de segredo.",
    ),
    (
        "18. Gestão de Conhecimento",
        "Cadastro inicial de documentos na Central.",
        "Editor rico, versionamento, requisitos vinculados, wiki, busca, anexos e integração real com Confluence.",
        "Criar documents, document_versions e links; editor Markdown/rich text; sync opcional via API do Confluence.",
        "Criar/editar/versionar; restaurar versão; vincular requisito; pesquisar; validar sincronização e conflitos.",
    ),
    (
        "19. DevOps",
        "Não há integração operacional; repositório público já definido para o portfólio.",
        "GitHub/GitLab/Bitbucket, commits, PRs, branches, deployments, pipelines e CI/CD vinculados.",
        "Conectar primeiro o GitHub; usar GitHub App/OAuth ou token de escopo mínimo; webhooks assinados; tabelas dev_links e deployments.",
        "Vincular commit/PR por chave; validar assinatura do webhook; atualizar deployment; testar falhas, retries e revogação.",
    ),
    (
        "20. Testes",
        "Cadastro inicial de casos de teste na Central.",
        "Planos, execuções, passos estruturados, evidências, cobertura e integrações Xray/Zephyr.",
        "Criar test_cases, test_steps, test_plans, executions e evidence; vincular issues; adaptadores opcionais para Xray/Zephyr.",
        "Executar caso aprovado/reprovado; anexar evidência; calcular cobertura; validar importação/exportação.",
    ),
    (
        "21. Integrações",
        "Catálogo inicial configurável para GitHub, GitLab, Slack, Teams e Confluence; ainda sem troca real de dados.",
        "Conectores reais para mensageria, e-mail, workspace, CI, CRM e BI, com credenciais e sincronização.",
        "Criar framework de conectores, cofre de segredos, health checks, logs e filas; implementar por prioridade de uso da equipe.",
        "Conectar sandbox de cada serviço; testar autenticação, envio/recebimento, retry, revogação e isolamento de segredos.",
    ),
    (
        "22. APIs",
        "API interna para estado completo, health check e restauração.",
        "REST API versionada por recurso, autenticação, paginação, filtros, rate limit, webhooks, documentação OpenAPI e SDK.",
        "Criar /api/v1; DTOs e validação; OpenAPI; tokens de serviço; idempotency keys; webhooks assinados.",
        "Contract tests; autenticação; paginação; limites; compatibilidade de versão; coleção Postman/OpenAPI validada.",
    ),
    (
        "23. Marketplace",
        "Não implementado e não é prioridade inicial para ferramenta interna.",
        "Arquitetura de plugins e módulos de financeiro, riscos, Gantt, PMO, CRM, RH, BI e aprovações.",
        "Após estabilizar a API, definir manifest de plugin, permissões e sandbox; começar com módulos internos aprovados.",
        "Instalar/desinstalar plugin sem corromper dados; validar permissões, compatibilidade e isolamento.",
    ),
    (
        "24. Recursos Corporativos",
        "Parte do planejamento cruzado pode ser construída sobre projetos e equipes; recursos enterprise ainda ausentes.",
        "Capacity Planning, programas, portfólios, sandboxes, auditoria avançada, data residency e insights executivos.",
        "Implementar somente após núcleo, segurança e relatórios; criar entidades program/portfolio e ambientes isolados.",
        "Validar consolidação entre projetos, capacidade, segregação de sandbox, auditoria e recuperação de dados.",
    ),
    (
        "25. Funcionalidades de IA — última etapa",
        "Não implementado por decisão de projeto.",
        "Resumo, geração de descrição, subtarefas, automações, busca semântica, resumo de comentários e assistência de escrita.",
        "Somente após segurança e governança: escolher provedor; política de dados; camada de redaction; prompts versionados; limites e logs.",
        "Avaliação humana com conjunto de casos; privacidade; precisão; custo; latência; prompt injection; opção de desativação.",
    ),
]


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(30)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(32, 33, 36)
for style_name, size, color in [
    ("Heading 1", 20, (103, 87, 217)),
    ("Heading 2", 15, (32, 33, 36)),
    ("Heading 3", 12, (52, 120, 201)),
]:
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.size = Pt(size)
    styles[style_name].font.bold = True
    styles[style_name].font.color.rgb = RGBColor(*color)
    styles[style_name].paragraph_format.space_before = Pt(10)
    styles[style_name].paragraph_format.space_after = Pt(5)

if "Callout" not in styles:
    callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
else:
    callout = styles["Callout"]
callout.font.name = "Aptos"
callout.font.size = Pt(10)
callout.font.color.rgb = RGBColor(73, 68, 121)
callout.paragraph_format.left_indent = Cm(0.6)
callout.paragraph_format.right_indent = Cm(0.6)
callout.paragraph_format.space_before = Pt(8)
callout.paragraph_format.space_after = Pt(8)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Plano para completar os 25 módulos")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Orbit Projects · Ferramenta interna de gestão de projetos")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(103, 87, 217)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Documento de implementação e validação · {date.today().strftime('%d/%m/%Y')}")

p = doc.add_paragraph(style="Callout")
p.add_run("Decisões já tomadas: ").bold = True
p.add_run("o sistema será uma ferramenta interna de equipe; PostgreSQL e a fundação técnica serão implementados; o repositório público será usado como portfólio; funcionalidades de IA ficarão por último.")

doc.add_heading("1. Objetivo do documento", level=1)
doc.add_paragraph(
    "Este plano transforma a lista de 25 grupos de funcionalidades em uma sequência executável. "
    "Para cada grupo, registra o estado atual, as lacunas, as ações de implementação e os critérios mínimos de validação. "
    "O princípio central é não considerar uma funcionalidade concluída apenas porque existe um botão: ela precisa persistir dados, respeitar permissões, ser testada e funcionar após recarregar o sistema."
)

doc.add_heading("2. Situação atual", level=1)
add_bullets(doc, [
    "Aplicação em JavaScript e Node.js, sem dependências obrigatórias de runtime.",
    "Interface SPA com projetos, tarefas, backlog, boards, sprints, roadmap, relatórios básicos, equipe e worklogs.",
    "Persistência atual em arquivo JSON; adequada para demonstração, mas insuficiente para uso simultâneo por uma equipe.",
    "Proteção inicial contra sobrescrita concorrente e migração automática do estado da interface.",
    "Recursos recentes: projetos híbridos, templates iniciais, clonagem, subtarefas, filtros salvos e Central da equipe.",
    "Ainda não existe autenticação real, banco relacional, anexos, filas, notificações ou integrações operacionais.",
])

doc.add_heading("3. Repositório GitHub e entrega contínua", level=1)
p = doc.add_paragraph()
p.add_run("Repositório público informado: ").bold = True
add_hyperlink(p, "araujofran/sistemaGestaoPro", REPO_URL)
doc.add_paragraph(
    "A pasta local atual não foi reconhecida pelo Git como uma worktree válida durante a elaboração deste documento. "
    "Antes de commitar, é necessário validar se o diretório .git está íntegro, inicializar o repositório local quando apropriado e configurar o remote origin sem sobrescrever histórico remoto."
)
doc.add_heading("Procedimento seguro", level=2)
add_numbered(doc, [
    "Consultar o conteúdo atual do repositório remoto e identificar a branch principal.",
    "Validar ou inicializar o repositório local; configurar origin para o endereço informado.",
    "Criar .env.example e garantir que .env, tokens, senhas, banco e anexos estejam no .gitignore.",
    "Executar testes e revisão de segredos antes do commit.",
    "Criar commits pequenos por módulo e enviar para uma branch de desenvolvimento.",
    "Usar Pull Request, mesmo em portfólio individual, para documentar decisões e manter histórico profissional.",
    "Configurar GitHub Actions para validação de sintaxe, testes, migrações e build.",
])

doc.add_heading("4. PostgreSQL e fundação técnica", level=1)
doc.add_paragraph(
    "PostgreSQL deve ser a primeira grande implementação. O arquivo JSON continuará apenas como fonte de migração e backup inicial. "
    "Para evitar uma reescrita arriscada, a migração deve ser incremental: criar banco e migrations, migrar leitura, migrar escrita, comparar resultados e somente então desativar o JSON."
)
add_numbered(doc, [
    "Adicionar configuração por variáveis de ambiente: DATABASE_URL, SESSION_SECRET, APP_URL e parâmetros de integração.",
    "Criar migrations versionadas para users, projects, issues, sprints, statuses, comments, worklogs, releases e activity_log.",
    "Criar constraints, índices, chaves estrangeiras e transações para exclusões e clonagens.",
    "Criar script idempotente para importar data/state.json para PostgreSQL.",
    "Separar backend em rotas, serviços, repositórios e validações.",
    "Substituir PUT do estado inteiro por endpoints REST granulares e controle otimista de versão.",
    "Adicionar seed de desenvolvimento, backup, restore e teste automático das migrations.",
])

doc.add_heading("Critérios de aceite do PostgreSQL", level=2)
add_bullets(doc, [
    "Todos os dados atuais são migrados sem perda de IDs ou relacionamentos.",
    "Duas pessoas podem editar itens diferentes simultaneamente sem sobrescrita.",
    "Operações compostas são atômicas e fazem rollback em caso de erro.",
    "Backup e restore são testados em ambiente separado.",
    "Nenhuma senha ou credencial aparece no código ou repositório.",
])

doc.add_heading("5. Método de implementação e validação", level=1)
add_numbered(doc, [
    "Definir história e critérios de aceite antes de codificar.",
    "Alterar schema por migration, nunca manualmente em produção.",
    "Implementar backend, autorização e validação antes de finalizar a interface.",
    "Criar testes unitários, de API e do fluxo principal no navegador.",
    "Validar com dados de exemplo e com um usuário sem permissão.",
    "Executar revisão de segurança, acessibilidade e responsividade.",
    "Registrar evidências e atualizar a documentação.",
    "Entregar por Pull Request e liberar primeiro em homologação.",
])

doc.add_heading("Definition of Done comum", level=2)
add_bullets(doc, [
    "Persistência e migrations concluídas.",
    "Permissões aplicadas no servidor, não somente ocultadas na interface.",
    "Mensagens de erro e confirmações de ações destrutivas.",
    "Testes automatizados e validação manual documentada.",
    "Sem segredos no Git e com logs de auditoria quando aplicável.",
    "Funciona após recarregar, em diferentes usuários e em telas menores.",
])

doc.add_page_break()
doc.add_heading("6. Plano dos 25 módulos", level=1)
doc.add_paragraph("A ordem numérica abaixo preserva a lista original. A ordem real de execução está na seção 7.")

for title_text, current, gaps, actions, validation in modules:
    doc.add_heading(title_text, level=2)
    add_status_table(doc, current, gaps, actions, validation)

doc.add_page_break()
doc.add_heading("7. Sequência recomendada de entrega", level=1)
phases = [
    ("Fase 0 — Fundação", "GitHub, CI, PostgreSQL, migrations, configuração, logs, testes básicos e migração do JSON."),
    ("Fase 1 — Núcleo operacional", "Projetos, tarefas, hierarquia, campos, backlog, Scrum, Kanban, boards, workflows e releases."),
    ("Fase 2 — Gestão e análise", "Roadmaps, dependências, tempo, relatórios, dashboards, busca avançada e filtros."),
    ("Fase 3 — Pessoas e segurança", "Login, usuários, grupos, times, papéis, permissões, auditoria, SSO e notificações."),
    ("Fase 4 — Colaboração e conhecimento", "Menções, watchers, anexos, wiki, requisitos, documentação e integração Confluence."),
    ("Fase 5 — DevOps, testes e integrações", "GitHub primeiro; depois mensageria, CI/CD, Xray/Zephyr e conectores realmente usados."),
    ("Fase 6 — APIs e extensibilidade", "REST v1, OpenAPI, webhooks, SDK e arquitetura de plugins internos."),
    ("Fase 7 — Corporativo", "Portfólios, programas, capacidade cruzada, sandboxes, auditoria avançada e insights."),
    ("Fase 8 — Inteligência Artificial", "Somente após segurança, dados, permissões e governança estarem estáveis."),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
set_cell_text(table.rows[0].cells[0], "Fase", bold=True, color=(255, 255, 255), size=10)
set_cell_text(table.rows[0].cells[1], "Resultado esperado", bold=True, color=(255, 255, 255), size=10)
shade(table.rows[0].cells[0], "6757D9")
shade(table.rows[0].cells[1], "6757D9")
for phase, result in phases:
    cells = table.add_row().cells
    set_cell_text(cells[0], phase, bold=True, size=9)
    set_cell_text(cells[1], result, size=9)

doc.add_heading("8. Validação com a equipe", level=1)
add_numbered(doc, [
    "Escolher dois usuários-piloto: um administrador e um membro comum.",
    "Criar um projeto real pequeno e executar uma sprint completa.",
    "Registrar bugs encontrados como issues do próprio sistema.",
    "Validar permissões e operações simultâneas em dois navegadores.",
    "Revisar métricas e relatórios com números calculados manualmente.",
    "Testar backup/restore antes de ampliar o uso.",
    "Liberar cada fase somente após aceite dos usuários-piloto.",
])

doc.add_heading("9. Próximas ações objetivas", level=1)
checklist = [
    "Conectar corretamente a pasta local ao repositório GitHub informado.",
    "Criar branch de desenvolvimento e primeiro commit organizado.",
    "Adicionar .env.example e política de segredos.",
    "Instalar/configurar PostgreSQL para desenvolvimento.",
    "Criar migrations e script de importação do JSON.",
    "Migrar projetos, tarefas, sprints, comentários, releases e worklogs.",
    "Criar autenticação e papéis para administrador, membro e leitor.",
    "Prosseguir pelos módulos na ordem das fases deste documento.",
    "Deixar IA para a última fase, conforme decisão registrada.",
]
for item in checklist:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("☐ " + item)
    run.font.size = Pt(10)

doc.add_heading("10. Conclusão", level=1)
doc.add_paragraph(
    "O projeto já possui uma boa base visual e operacional para portfólio. Para se tornar uma ferramenta interna confiável, "
    "o próximo marco não é adicionar mais telas isoladas: é consolidar PostgreSQL, autenticação, permissões, API modular, testes e CI. "
    "A partir dessa fundação, os 25 módulos podem ser entregues e validados gradualmente, mantendo o sistema utilizável durante toda a evolução."
)

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Orbit Projects · Plano de implementação dos 25 módulos")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(112, 116, 124)

doc.save(OUTPUT)
print(OUTPUT)
